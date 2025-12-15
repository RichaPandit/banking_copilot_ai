from docx import Document
from fastapi import APIRouter, Header, HTTPException
from rag_logic.rag_highlights import generate_rag_highlights
from rag_logic.risk_scoring import compute_risk_score
import pandas as pd
import datetime
import os
import requests
from typing import Dict, Optional

router = APIRouter()

# Use a writable base path on App Service Linux
BASE_DIR = os.environ.get("APP_BASE_DIR", "/home/site/wwwroot")
REPORT_DIR = os.path.join(BASE_DIR, "reports", "generated_reports")

# URL-building inputs
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "").strip().rstrip("/")
WEBSITE_HOSTNAME = os.getenv("WEBSITE_HOSTNAME", "").strip()

TEAMS_WORKFLOW_WEBHOOK_URL: str = os.getenv("TEAMS_WORKFLOW_WEBHOOK_URL", "").strip()

def _to_web_path(*parts: str) -> str:
    rel = os.path.join(*parts)
    return rel.replace(os.sep, "/").lstrip("/")

def _public_base_url() -> str:
    if PUBLIC_BASE_URL:
        return PUBLIC_BASE_URL
    if WEBSITE_HOSTNAME:
        return f"https://{WEBSITE_HOSTNAME}"
    return ""

def _make_public_url(rel_web_path: Optional[str]) -> Optional[str]:
    if not rel_web_path:
        return None
    base = _public_base_url()
    return f"{base}/{rel_web_path.lstrip('/')}" if base else None

def escalate_alert_internal(
    company_id: str,
    x_agent_id: str,
    risk_score: Optional[float] = None,
    risk_rating: Optional[str] = None,
    extra: Optional[Dict] = None,
) -> Dict[str, str]:

    if not x_agent_id or not x_agent_id.startswith("agent-"):
        raise HTTPException(status_code=401, detail="Invalid or missing Agent ID")

    if not TEAMS_WORKFLOW_WEBHOOK_URL:
        raise HTTPException(status_code=500, detail="Teams webhook URL not configured")

    # Minimal Adaptive Card (you can customize this in the Adaptive Cards Designer)
    # https://adaptivecards.io/designer/
    card = {
        "type": "AdaptiveCard",
        "version": "1.5",
        "body": [
            {"type": "TextBlock", "size": "Large", "weight": "Bolder",
             "text": f"🚨 Risk Escalation: {company_id}"},
            {"type": "TextBlock",
             "text": "Triggered by Credit Risk MCP (Copilot Studio)"},
            {"type": "FactSet", "facts": [
                {"title": "Company ID", "value": company_id},
                {"title": "Agent", "value": x_agent_id},
                {"title": "Risk Score", "value": "" if risk_score is None else f"{risk_score:.4f}"},
                {"title": "Risk Rating", "value": risk_rating or ""},
                {"title": "Timestamp (UTC)", "value": datetime.datetime.utcnow().isoformat() + "Z"},
            ]},
        ],
        # optional actions (buttons)
        "actions": [
            {
                "type": "Action.OpenUrl",
                "title": "Open Generated Report",
                # If you return absolute URLs in generate_report, set that here:
                "url": extra.get("report_url") if extra else
                "https://contoso.example/reports"
            }
        ]
    }

    try:
        resp = requests.post(
            TEAMS_WORKFLOW_WEBHOOK_URL,
            json=card,
            timeout=10
        )
        ok = 200 <= resp.status_code < 300
        return {
            "status": "alert_escalated" if ok else "failed",
            "code": str(resp.status_code),
            "message": "" if ok else resp.text[:300]
        }
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=502, detail=f"Webhook post failed: {e}")

# ----------------------------------------------
# Word Report Generation
# ----------------------------------------------
def create_word_report(company_name, financials, exposure, covenants, ews, risk_score, risk_rating, rag_highlights):

    # Validate data presence to avoid iloc errors
    if financials.empty or exposure.empty or covenants.empty:
        raise HTTPException(status_code=400, detail="Insufficient data to generate report")
 
    # Ensure we pick the latest financials
    if "year" in financials.columns:
        financials = financials.sort_values("year")

    doc = Document()
    doc.add_heading(f"{company_name} - Quaterly Risk Review", level=0)

    # Financial Summary
    doc.add_heading("Financial Summary:", level=1)
    latest_fin = financials.iloc[-1]
    doc.add_paragraph(f"Revenue: {latest_fin['revenue']}")
    doc.add_paragraph(f"EBITDA: {latest_fin['ebitda']}")
    doc.add_paragraph(f"Net Income: {latest_fin['net_income']}")

    # Loan Exposure
    doc.add_heading("Loan Exposure:", level=1)
    exp = exposure.iloc[0]
    doc.add_paragraph(f"Sanctioned Limit: {exp['sanctioned_limit']}")
    doc.add_paragraph(f"Utilized Amount: {exp['utilized_amount']}")
    doc.add_paragraph(f"Overdue Amount: {exp['overdue_amount']}")

    # Covenant Compliance
    doc.add_heading("Covenant Compliance:", level=1)
    cov = covenants.iloc[0]
    doc.add_paragraph(f"DSCR: {cov['dscr']}")
    doc.add_paragraph(f"Interest Coverage: {cov['interest_coverage']}")
    doc.add_paragraph(f"Current Ratio: {cov['current_ratio']}")

    # Early Warning Signals
    doc.add_heading("Early Warning Signals:", level=1)
    for idx, row in ews.iterrows():
        doc.add_paragraph(f"{row['event_date']}: {row['event']}")
    
    # Risk Summary
    doc.add_heading("Risk Summary:", level=1)
    doc.add_paragraph(f"Risk Score: {risk_score:.2f}")
    doc.add_paragraph(f"Risk Rating: {risk_rating}")

    # Key Highlights (RAG)
    doc.add_heading("Key Highlights:", level=1)
    for h in rag_highlights:
        doc.add_paragraph(f"- {h}")

    # ----- Save under wwwroot so it's web-accessible -----
    os.makedirs(REPORT_DIR, exist_ok=True)
    report_name = f"{company_name.replace(' ','_')}_Risk_Report_{datetime.date.today()}.docx"
    # Absolute filesystem path
    abs_word_path = os.path.join(REPORT_DIR, report_name)
    # Relative web path (used for URL building)
    rel_web_word_path = _to_web_path("reports", "generated_reports", report_name)

    doc.save(abs_word_path)

    # Convert to PDF (best-effort, likely None on Linux)
    abs_pdf_path = abs_word_path.replace(".docx", ".pdf")
    rel_web_pdf_path = rel_web_word_path.replace(".docx", ".pdf")
    try:
        from docx2pdf import convert  # requires Word on Windows/macOS
        convert(abs_word_path, abs_pdf_path)
        pdf_exists = True
    except Exception:
        pdf_exists = False
        abs_pdf_path = None
        rel_web_pdf_path = None

    # Build URLs (absolute)
    word_url = _make_public_url(rel_web_word_path)
    pdf_url  = _make_public_url(rel_web_pdf_path) if pdf_exists else None

    return {
        # Return relative web paths for compatibility…
        "word": rel_web_word_path,              # 'reports/generated_reports/…docx'
        "pdf": rel_web_pdf_path,                # or None
        # …and public URLs for direct clicking
        "word_url": word_url,                   # 'https://…/reports/generated_reports/…docx'
        "pdf_url": pdf_url,                     # or None
        "rag_highlights": rag_highlights
    }

def generate_report_internal(
    company_id: str,
    x_agent_id: str,
    companies: pd.DataFrame,
    financials: pd.DataFrame,
    exposure: pd.DataFrame,
    covenants: pd.DataFrame,
    ews: pd.DataFrame
) -> Dict[str, str]:
    if not x_agent_id or not x_agent_id.startswith("agent-"):
        raise HTTPException(status_code=401, detail="Invalid or missing Agent ID")

    # Validate company
    company_info = companies[companies["company_id"] == company_id]
    if company_info.empty:
        raise HTTPException(status_code=404, detail="Company not found")
    company_name = company_info["company_name"].values[0]

    # Extract data
    fin = financials[financials["company_id"] == company_id]
    exp = exposure[exposure["company_id"] == company_id]
    cov = covenants[covenants["company_id"] == company_id]
    ews_events = ews[ews["company_id"] == company_id]

    # Compute risk score
    risk_score, risk_rating = compute_risk_score(fin, exp, cov, ews_events)

    # Generate reports
    rag_highlights = generate_rag_highlights(fin, cov, exp, ews_events)
    report_paths = create_word_report(company_name, fin, exp, cov, ews_events, risk_score, risk_rating, rag_highlights)

    return {
        "status": "report_generated",
        "company": company_name,
        "risk_score": risk_score,
        "risk_rating": risk_rating,
        "word_report": report_paths["word"],              # relative web path
        "word_report_url": report_paths["word_url"],      # absolute URL
        "pdf_report": report_paths["pdf"],                # relative web path or None
        "pdf_report_url": report_paths["pdf_url"],        # absolute URL or None
        "rag_highlights": report_paths["rag_highlights"]
    }